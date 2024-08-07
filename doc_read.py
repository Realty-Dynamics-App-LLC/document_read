from docx import Document
import json
import os
import docx
import docx2pdf
import camelot
import pdfplumber

extension = ".docx" 
file_list = []

# Get all filenames in the current directory
filenames = os.listdir()

# Iterate through filenames
for filename in filenames:
  if os.path.isfile(filename):       # Check if it's a file (not a directory)
    if filename.endswith(extension):
      file_list.append(filename)

def read_doc(document):
    # Convert DOCX to PDF
    pdf_file = "temp.pdf"
    if not os.path.exists(pdf_file):
        docx2pdf.convert(docx_file, pdf_file)

    # Extract tables from PDF using pdfplumber
    tables_data = []
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                # Format the table as a string
                table_str = ""
                for row in table:
                    table_str += "\t".join(str(cell) for cell in row) + "\n"
                tables_data.append(table_str)    

    # Extract text content and tables from DOCX file
    content = []

    tables_index = 0

    for element in document.element.body:
        if element.tag.endswith('p'):  # Paragraph
            para = docx.text.paragraph.Paragraph(element, document)
            content.append(para.text)
        elif element.tag.endswith('tbl'):  # Table
            if tables_index < len(tables_data):
                content.append(tables_data[tables_index])  # Insert the formatted table data
                tables_index += 1
    # Combine the content into a single string
    final_content = "\n".join(content)
    return final_content

def text_read(document):
   #Accumulate content from all paragraphs
   content = ""
   for paragraph in document.paragraphs:
    content += paragraph.text + "\n"  # Add paragraph text with newline

   return content

def create_link(new_filename):  #creates article link
    gen_link = "https://library.municode.com/fl/brevard_county/codes/code_of_ordinances?nodeId=COORBRCOFLVOII_CH62LADERE_ARTVIZORE"   #chapter link
    #article = "ARTIINPUME" 
    filename_split = new_filename.split('_')

    for i in filename_split:
      if i in ['AND','OR','FOR','OF']:
        filename_split.remove(i)

    article_no = filename_split[1]

    article_name = 'ART' + article_no  # can be + '_1-2_' , "div"
    n = len(filename_split)
    for i in range(2,n):
      article_name += filename_split[i][0:2]

    final_link = gen_link + '_' + article_name
    return final_link
   
for i in file_list:
  filename = i
  county_name = ''   #enter respective county name

  document = Document(filename)
  docx_file = filename

  #skipping extension to get only the filename
  fn = len(filename)
  skip_extension = fn - 4
  filename = i[0:skip_extension]

  new_filename = filename.replace('.',"")   #replacing any other periods in filename to avoid wrong extension errors
  link = create_link(new_filename)          #replace with document's link from municode 

  if len(document.tables) > 0:   #if tables present in document
    final_content = read_doc(document)

  else:
     final_content = text_read(document)

  # Create an empty dictionary
  data = {
    "link": link,
    "content": final_content  # Initialize empty string for content
  }

  # Write the data to a JSON file
  json_filename = f"{county_name}_{new_filename}.json"

  with open(json_filename, 'w')  as outfile:
    json.dump(data, outfile, indent = 4)

  print("Text content successfully written")
