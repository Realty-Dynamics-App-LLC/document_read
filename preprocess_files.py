import argparse
from docx import Document
import re
from multiprocessing import Pool
import os
import json

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            try:
                row_text = '\t'.join(cell.text for cell in row.cells)
                full_text.append(row_text)
            except ValueError as e:
                print(f"Skipping problematic row in table in file {file_path}: {e}")
                continue

    return '\n'.join(full_text)

def clean_text(text):
    # Remove unwanted characters
    text = re.sub(r'\n+', '\n', text)  # Replace multiple newlines with a single newline
    text = re.sub(r'\s+', ' ', text)   # Replace multiple spaces with a single space
    text = text.strip()  # Remove leading and trailing whitespace
    return text

def segment_text(text, max_length=2048):
    words = text.split(' ')
    segments = []
    segment = []

    for word in words:
        if len(' '.join(segment + [word])) > max_length:
            segments.append(' '.join(segment))
            segment = [word]
        else:
            segment.append(word)
    segments.append(' '.join(segment))
    return segments

def preprocess_document(file_path):
    document_text = read_docx(file_path)
    cleaned_text = clean_text(document_text)
    text_segments = segment_text(cleaned_text)
    return text_segments

def process_documents_in_parallel(file_paths):
    with Pool(processes=os.cpu_count()) as pool:
        results = pool.map(preprocess_document, file_paths)
    return results

def save_preprocessed_text(file_path, segments):
    file_name = os.path.basename(file_path)
    file_link = file_name.replace('.docx', '')
    content = "\n".join(segments)

    output_dict = {
        "link": file_link,
        "content": content
    }

    output_file_path = os.path.splitext(file_path)[0] + '.txt'
    with open(output_file_path, 'w') as out_file:
        json.dump(output_dict, out_file, indent=4)

def main(input_dir):
    # List all .docx files in the input directory
    file_paths = [os.path.join(input_dir, f) for f in os.listdir(input_dir) if f.endswith('.docx')]
    
    # Process all documents in parallel
    preprocessed_texts = process_documents_in_parallel(file_paths)

    # Save each preprocessed document to its own output file
    for file_path, segments in zip(file_paths, preprocessed_texts):
        save_preprocessed_text(file_path, segments)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Preprocess .docx files for embeddings.')
    parser.add_argument('input_dir', type=str, help='Directory containing .docx files to preprocess')
    
    args = parser.parse_args()
    main(args.input_dir)
